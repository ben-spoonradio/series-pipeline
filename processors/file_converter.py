"""
File Converter Processor
Converts various file formats (docx, pdf, hwp, epub, rtf) to plain text
"""

from typing import Dict, Any, Optional
import os
import shutil
import subprocess
import logging

from processors.base_processor import BaseProcessor, ProcessorType

logger = logging.getLogger(__name__)


class FileConverter(BaseProcessor):
    """
    File format converter processor

    Converts multiple formats to plain text:
    - .docx (Microsoft Word)
    - .pdf (Adobe PDF)
    - .hwp (Hancom HWP)
    - .epub (E-book)
    - .rtf (Rich Text Format)
    - .txt (Plain text - passthrough with encoding handling)
    - .md (Markdown - passthrough)
    """

    def __init__(self):
        super().__init__(ProcessorType.RULE_BASED)
        self.supported_formats = ['.docx', '.pdf', '.hwp', '.epub', '.rtf', '.txt', '.md']

    def process(self, input_data: Dict[str, Any]) -> Dict[str, Any]:
        """
        Convert file to plain text

        Args:
            input_data: {
                'file_path': str (absolute path to file),
                'file_content': bytes (optional, file content as bytes)
            }

        Returns:
            {
                'output': str (extracted text),
                'format': str (original file extension),
                'metadata': {
                    'original_filename': str,
                    'file_size': int,
                    'conversion_method': str
                }
            }
        """
        file_path = input_data.get('file_path')
        file_content = input_data.get('file_content')

        if not file_path and not file_content:
            raise ValueError("Either file_path or file_content must be provided")

        # Determine file extension
        if file_path:
            ext = os.path.splitext(file_path)[1].lower()
            original_filename = os.path.basename(file_path)
        else:
            # If only content provided, try to infer from metadata
            ext = input_data.get('extension', '.txt').lower()
            original_filename = input_data.get('filename', 'uploaded_file' + ext)

        # Validate format support
        if ext not in self.supported_formats:
            raise ValueError(f"Unsupported file format: {ext}. Supported: {self.supported_formats}")

        # Extract text based on format
        text = self._extract_text(file_path, file_content, ext)

        return {
            'output': text,
            'format': ext,
            'metadata': {
                'original_filename': original_filename,
                'file_size': len(file_content) if file_content else os.path.getsize(file_path),
                'conversion_method': self._get_conversion_method(ext),
                'text_length': len(text)
            }
        }

    def validate(self, output_data: Dict[str, Any]) -> bool:
        """Validate conversion output"""
        text = output_data.get('output', '')

        # Check if text was extracted
        if not text or not text.strip():
            self.logger.warning("Conversion produced empty text")
            return False

        # Check minimum length (at least 10 characters)
        if len(text.strip()) < 10:
            self.logger.warning(f"Converted text too short: {len(text)} characters")
            return False

        return True

    def _extract_text(self, file_path: Optional[str], file_content: Optional[bytes], ext: str) -> str:
        """Route to appropriate extraction method"""
        if ext == '.docx':
            return self._extract_text_from_docx(file_path, file_content)
        elif ext == '.pdf':
            return self._extract_text_from_pdf(file_path, file_content)
        elif ext == '.hwp':
            return self._extract_text_from_hwp(file_path, file_content)
        elif ext == '.epub':
            return self._extract_text_from_epub(file_path, file_content)
        elif ext == '.rtf':
            return self._extract_text_from_rtf(file_path, file_content)
        elif ext in ['.txt', '.md']:
            return self._extract_text_from_txt(file_path, file_content)
        else:
            raise ValueError(f"Unsupported format: {ext}")

    def _extract_text_from_docx(self, file_path: Optional[str], file_content: Optional[bytes]) -> str:
        """Extract text from DOCX file"""
        try:
            from docx import Document

            if file_content:
                from io import BytesIO
                doc = Document(BytesIO(file_content))
            else:
                doc = Document(file_path)

            full_text = []
            for para in doc.paragraphs:
                full_text.append(para.text)

            return '\n'.join(full_text)
        except ImportError:
            raise ImportError("python-docx not installed. Run: pip install python-docx")
        except Exception as e:
            self.logger.error(f"Error reading docx: {e}")
            return ""

    def _extract_text_from_pdf(self, file_path: Optional[str], file_content: Optional[bytes]) -> str:
        """Extract text from PDF file"""
        try:
            from pypdf import PdfReader

            if file_content:
                from io import BytesIO
                reader = PdfReader(BytesIO(file_content))
            else:
                reader = PdfReader(file_path)

            text = ""
            for page in reader.pages:
                text += page.extract_text() + "\n"

            return text
        except ImportError:
            raise ImportError("pypdf not installed. Run: pip install pypdf")
        except Exception as e:
            self.logger.error(f"Error reading pdf: {e}")
            return ""

    def _extract_text_from_epub(self, file_path: Optional[str], file_content: Optional[bytes]) -> str:
        """Extract text from EPUB file"""
        try:
            from ebooklib import epub
            import ebooklib
            from bs4 import BeautifulSoup

            if file_content:
                from io import BytesIO
                book = epub.read_epub(BytesIO(file_content))
            else:
                book = epub.read_epub(file_path)

            text = []
            for item in book.get_items():
                if item.get_type() == ebooklib.ITEM_DOCUMENT:
                    soup = BeautifulSoup(item.get_content(), 'html.parser')
                    text.append(soup.get_text())

            return '\n'.join(text)
        except ImportError:
            raise ImportError("ebooklib or beautifulsoup4 not installed. Run: pip install ebooklib beautifulsoup4")
        except Exception as e:
            self.logger.error(f"Error reading epub: {e}")
            return ""

    def _extract_text_from_rtf(self, file_path: Optional[str], file_content: Optional[bytes]) -> str:
        """Extract text from RTF file"""
        try:
            from striprtf.striprtf import rtf_to_text

            if file_content:
                content = file_content.decode('utf-8')
            else:
                # Try UTF-8 first
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        content = f.read()
                except UnicodeDecodeError:
                    # Fallback to cp949 (common for Korean/Japanese)
                    with open(file_path, 'r', encoding='cp949') as f:
                        content = f.read()

            return rtf_to_text(content)
        except ImportError:
            raise ImportError("striprtf not installed. Run: pip install striprtf")
        except Exception as e:
            self.logger.error(f"Error reading rtf: {e}")
            return ""

    def _extract_text_from_hwp(self, file_path: Optional[str], file_content: Optional[bytes]) -> str:
        """
        Extract text from HWP file using hwp5txt command-line tool

        Note: Requires hwp5txt to be installed on system
        Installation: pip install hwp5 && pip install olefile
        """
        try:
            import tempfile

            # HWP extraction requires file path (hwp5txt CLI tool)
            if file_content:
                # Save to temporary file
                with tempfile.NamedTemporaryFile(suffix='.hwp', delete=False) as temp_file:
                    temp_file.write(file_content)
                    temp_path = temp_file.name
            else:
                # Copy to temp to avoid encoding issues
                temp_path = tempfile.mktemp(suffix='.hwp')
                shutil.copy(file_path, temp_path)

            try:
                # Run hwp5txt command (try full path first, then PATH)
                hwp5txt_paths = [
                    os.path.expanduser('~/.local/bin/hwp5txt'),
                    'hwp5txt'
                ]
                for hwp5txt_cmd in hwp5txt_paths:
                    try:
                        result = subprocess.check_output([hwp5txt_cmd, temp_path], stderr=subprocess.STDOUT)
                        return result.decode('utf-8')
                    except FileNotFoundError:
                        continue
                raise FileNotFoundError("hwp5txt not found")
            finally:
                # Clean up temp file
                if os.path.exists(temp_path):
                    os.remove(temp_path)

        except FileNotFoundError:
            raise ImportError("hwp5txt not installed. Run: pip install hwp5 olefile")
        except Exception as e:
            self.logger.error(f"Error reading hwp: {e}")
            return ""

    def _extract_text_from_txt(self, file_path: Optional[str], file_content: Optional[bytes]) -> str:
        """Extract text from TXT/MD file with encoding handling"""
        try:
            if file_content:
                # Try UTF-8 first
                try:
                    return file_content.decode('utf-8')
                except UnicodeDecodeError:
                    # Fallback to cp949 (common for Korean/Japanese)
                    return file_content.decode('cp949')
            else:
                # Try UTF-8 first
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        return f.read()
                except UnicodeDecodeError:
                    # Fallback to cp949
                    with open(file_path, 'r', encoding='cp949') as f:
                        return f.read()
        except Exception as e:
            self.logger.error(f"Error reading text file: {e}")
            return ""

    def _get_conversion_method(self, ext: str) -> str:
        """Get human-readable conversion method name"""
        methods = {
            '.docx': 'python-docx',
            '.pdf': 'pypdf',
            '.hwp': 'hwp5txt',
            '.epub': 'ebooklib + BeautifulSoup',
            '.rtf': 'striprtf',
            '.txt': 'direct read',
            '.md': 'direct read'
        }
        return methods.get(ext, 'unknown')
